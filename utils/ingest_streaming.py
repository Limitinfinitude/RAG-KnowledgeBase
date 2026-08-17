"""
大文件入库：分段读入、单层（medium）切分、分批 add_documents，降低峰值内存与单次 embedding 批量。
"""
from __future__ import annotations

import gc
import os
from typing import Generator, Iterable, List, Optional, Tuple

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from utils.document_parsers import detect_text_file_encoding  # noqa: F401 — 兼容旧引用 ins.detect_text_file_encoding
from utils.path_context import get_kb_dir
from utils.smart_chunker import CHINESE_SEPARATORS, SmartChunker
from utils.web_system_settings import get_merged_chunk_levels

# 超过此字节数走流式路径（约 9～10 万汉字 UTF-8 量级，可按机器内存再调）
STREAMING_MIN_BYTES = 280_000
# 每段目标字符数（单段内做 Recursive 切分，避免整书进 smart_chunk 多层级）
SEGMENT_CHAR_TARGET = 55_000
# 单次写入向量库的文档数上限（控制 embedding 批大小）
EMBED_ADD_BATCH_SIZE = 20
# 每累计若干次向量写入后落盘 FAISS，降低中途崩溃损失
STREAM_SAVE_EVERY_FLUSHES = 6


def should_use_streaming_ingest(file_size_bytes: int) -> bool:
    return int(file_size_bytes) >= STREAMING_MIN_BYTES


def doc_length_factor_from_filesize(file_size_bytes: int) -> float:
    sz = int(file_size_bytes)
    if sz < 120_000:
        return 0.9
    if sz > 900_000:
        return 1.15
    return 1.0


def read_text_head(path: str, encoding: str, max_chars: int = 16000) -> str:
    with open(path, "r", encoding=encoding, errors="replace") as f:
        return f.read(max_chars)


def _medium_splitter(doc_length_factor: float) -> RecursiveCharacterTextSplitter:
    merged = get_merged_chunk_levels()
    cfg = merged.get("medium") or {"chunk_size": 800, "chunk_overlap": 100}
    cs = int(cfg["chunk_size"] * doc_length_factor)
    co = int(cfg["chunk_overlap"] * doc_length_factor)
    cs = max(200, min(cs, 16000))
    co = max(0, min(co, max(cs - 1, 0)))
    return RecursiveCharacterTextSplitter(
        chunk_size=cs,
        chunk_overlap=co,
        separators=CHINESE_SEPARATORS,
        length_function=len,
    )


def split_segment_medium(
    text: str,
    source_file: str,
    file_type: str,
    segment_index: int,
    chunk_index_start: int,
    doc_length_factor: float,
) -> Tuple[List[Document], int]:
    text = (text or "").strip()
    if not text:
        return [], chunk_index_start
    splitter = _medium_splitter(doc_length_factor)
    base = Document(
        page_content=text,
        metadata={
            "source_file": source_file,
            "file_type": file_type,
            "chunk_level": "medium",
            "stream_segment": segment_index,
            "doc_type": "long_document",
        },
    )
    parts = splitter.split_documents([base])
    out: List[Document] = []
    for i, ch in enumerate(parts):
        ch.metadata["chunk_index"] = chunk_index_start + i
        ch.metadata["source_file"] = source_file
        ch.metadata["file_type"] = file_type
        ch.metadata["chunk_level"] = "medium"
        ch.metadata["stream_segment"] = segment_index
        ch.metadata.setdefault("doc_type", "long_document")
        out.append(ch)
    return out, chunk_index_start + len(out)


def make_rule_summary_from_sample(sample_text: str, source_file: str, file_type: str) -> Optional[Document]:
    if not sample_text or len(sample_text.strip()) < 80:
        return None
    sc = SmartChunker()
    return sc._create_summary_chunk(sample_text, source_file, file_type, use_llm=False)


def iter_segments_text_file(path: str, encoding: str) -> Generator[str, None, None]:
    buf: List[str] = []
    size = 0
    target = SEGMENT_CHAR_TARGET
    with open(path, "r", encoding=encoding, errors="replace") as f:
        for line in f:
            buf.append(line)
            size += len(line)
            if size >= target and (line.strip() == "" or size >= int(target * 1.3)):
                yield "".join(buf)
                buf = []
                size = 0
        if buf:
            yield "".join(buf)


def _pdf_reader_class():
    try:
        from pypdf import PdfReader

        return PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader

            return PdfReader
        except ImportError:
            return None


def pdf_page_count(path: str) -> int:
    PdfReader = _pdf_reader_class()
    if PdfReader is None:
        return 0
    try:
        return len(PdfReader(path).pages)
    except Exception:
        return 0


def pdf_has_extractable_text(path: str, max_pages_sample: int = 5) -> bool:
    PdfReader = _pdf_reader_class()
    if PdfReader is None:
        return False
    try:
        reader = PdfReader(path)
        for i, page in enumerate(reader.pages):
            if i >= max_pages_sample:
                break
            t = page.extract_text() or ""
            if t.strip():
                return True
    except Exception:
        return False
    return False


def iter_segments_pdf_text(path: str) -> Generator[str, None, None]:
    PdfReader = _pdf_reader_class()
    if PdfReader is None:
        return
    reader = PdfReader(path)
    buf: List[str] = []
    size = 0
    target = SEGMENT_CHAR_TARGET
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        buf.append(t)
        size += len(t)
        if size >= target:
            yield "\n\n".join(buf)
            buf = []
            size = 0
    if buf:
        yield "\n\n".join(buf)


def read_pdf_text_sample(path: str, max_chars: int = 16000) -> str:
    PdfReader = _pdf_reader_class()
    if PdfReader is None:
        return ""
    out: List[str] = []
    n = 0
    try:
        reader = PdfReader(path)
        for page in reader.pages:
            t = page.extract_text() or ""
            out.append(t)
            n += len(t)
            if n >= max_chars:
                break
    except Exception:
        return ""
    return "\n\n".join(out)[:max_chars]


def iter_segments_docx(path: str) -> Generator[str, None, None]:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    buf: List[str] = []
    size = 0
    target = SEGMENT_CHAR_TARGET

    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if not t:
            continue
        buf.append(t)
        size += len(t) + 2
        if size >= target:
            yield "\n\n".join(buf)
            buf = []
            size = 0
    for table in doc.tables:
        rows_txt: List[str] = []
        for row in table.rows:
            cells = " | ".join((c.text or "").strip() for c in row.cells)
            if cells.strip():
                rows_txt.append(cells)
        if rows_txt:
            block = "\n".join(rows_txt)
            buf.append(block)
            size += len(block) + 2
            if size >= target:
                yield "\n\n".join(buf)
                buf = []
                size = 0
    if buf:
        yield "\n\n".join(buf)


def read_docx_sample(path: str, max_chars: int = 16000) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    parts: List[str] = []
    n = 0
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if not t:
            continue
        parts.append(t)
        n += len(t) + 2
        if n >= max_chars:
            break
    return "\n\n".join(parts)[:max_chars]


def iter_segments_pdf_ocr(path: str, dpi: int = 200) -> Generator[str, None, None]:
    """扫描版 PDF：按页渲染两层 OCR（本地 → 云端回退），按字数合并为段，避免一次性载入全书图像。"""
    from pdf2image import convert_from_path

    from utils.document_parsers import _ocr_image

    n = pdf_page_count(path)
    if n <= 0:
        return
    buf: List[str] = []
    size = 0
    for i in range(n):
        imgs = convert_from_path(path, first_page=i + 1, last_page=i + 1, dpi=dpi, fmt="png")
        try:
            img = imgs[0]
            page_text, _engine = _ocr_image(img)
        finally:
            del imgs
        block = "第%d页：\n%s" % (i + 1, page_text)
        buf.append(block)
        size += len(block)
        if size >= SEGMENT_CHAR_TARGET:
            yield "\n\n".join(buf)
            buf = []
            size = 0
    if buf:
        yield "\n\n".join(buf)


def run_streaming_ingest(
    segment_iter: Iterable[str],
    source_file: str,
    file_type: str,
    vector_db,
    file_size_bytes: int,
    summary_doc: Optional[Document],
) -> int:
    """
    消费文本段迭代器：切 medium chunk、分批 add_documents、周期性 save_local。
    返回写入的 chunk 条数（含可选 1 条 summary）。

    全程持有 FAISS 写锁：大文件入库可达数分钟，期间同目录的删除/重置/另一入库
    必须等待，否则周期性 save_local 会互相覆盖索引。
    """
    from utils.faiss_write_lock import faiss_write_lock

    with faiss_write_lock():
        return _run_streaming_ingest_unlocked(
            segment_iter, source_file, file_type, vector_db, file_size_bytes, summary_doc
        )


def _run_streaming_ingest_unlocked(
    segment_iter: Iterable[str],
    source_file: str,
    file_type: str,
    vector_db,
    file_size_bytes: int,
    summary_doc: Optional[Document],
) -> int:
    index_dir = os.path.join(get_kb_dir(), "faiss_index")
    os.makedirs(index_dir, exist_ok=True)
    factor = doc_length_factor_from_filesize(file_size_bytes)
    doc_length_factor = factor

    total = 0
    chunk_i = 0
    seg_i = 0
    pending: List[Document] = []
    flushes_since_save = 0

    def flush_batch(docs: List[Document]) -> None:
        nonlocal total, flushes_since_save
        if not docs:
            return
        vector_db.add_documents(docs)
        total += len(docs)
        flushes_since_save += 1
        if flushes_since_save >= STREAM_SAVE_EVERY_FLUSHES:
            vector_db.save_local(index_dir)
            flushes_since_save = 0
        gc.collect()

    if summary_doc is not None:
        flush_batch([summary_doc])

    for segment in segment_iter:
        chunks, chunk_i = split_segment_medium(
            segment, source_file, file_type, seg_i, chunk_i, doc_length_factor
        )
        seg_i += 1
        for ch in chunks:
            pending.append(ch)
            if len(pending) >= EMBED_ADD_BATCH_SIZE:
                batch = pending[:EMBED_ADD_BATCH_SIZE]
                del pending[:EMBED_ADD_BATCH_SIZE]
                flush_batch(batch)

    if pending:
        flush_batch(pending)

    vector_db.save_local(index_dir)
    return total


__all__ = [
    "STREAMING_MIN_BYTES",
    "EMBED_ADD_BATCH_SIZE",
    "should_use_streaming_ingest",
    "detect_text_file_encoding",
    "read_text_head",
    "iter_segments_text_file",
    "iter_segments_pdf_text",
    "iter_segments_docx",
    "read_pdf_text_sample",
    "read_docx_sample",
    "pdf_has_extractable_text",
    "pdf_page_count",
    "iter_segments_pdf_ocr",
    "make_rule_summary_from_sample",
    "run_streaming_ingest",
    "doc_length_factor_from_filesize",
]
