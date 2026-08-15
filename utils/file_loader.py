# utils/file_loader.py
import os
import tempfile
import traceback
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
from config import TESSERACT_CMD
from utils.path_context import get_kb_dir
from utils.document_preview import persist_original_from_temp
from utils.metadata_manager import MAX_FILE_SIZE_BYTES, add_document_metadata, update_chunks_count
from utils.logger import log_file_upload, log_error

pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

__all__ = ["ingest_file", "MAX_FILE_SIZE_BYTES"]


def _finalize_ingest_metadata(
    uploaded_file,
    file_ext: str,
    category: str,
    description: str,
    chunks_count: int,
) -> None:
    if chunks_count <= 0:
        return
    file_size = len(uploaded_file.getbuffer())
    add_document_metadata(
        file_name=uploaded_file.name,
        file_size=file_size,
        file_type=file_ext.lstrip("."),
        category=category,
        description=description,
    )
    update_chunks_count(uploaded_file.name, chunks_count)
    try:
        log_file_upload(
            file_name=uploaded_file.name,
            file_size=file_size,
            chunks=chunks_count,
            category=category,
        )
    except Exception as e:
        print(f"记录上传日志失败: {e}")


def ingest_file(uploaded_file, vector_db, category: str = "默认知识库", description: str = ""):
    """
    处理上传的文件并入库。
    大文件走流式：分段读入、单层 medium 切分、分批写入向量库，降低内存峰值。
    """
    from utils import ingest_streaming as ins

    temp_path = None
    try:
        file_size = len(uploaded_file.getbuffer())
        if file_size > MAX_FILE_SIZE_BYTES:
            raise ValueError(
                f"文件大小 {file_size / (1024*1024):.2f}MB 超过限制 {MAX_FILE_SIZE_BYTES / (1024*1024)}MB"
            )

        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext, mode="wb") as tmp:
            tmp.write(uploaded_file.getbuffer())
            temp_path = tmp.name

        persist_original_from_temp(temp_path, uploaded_file.name)

        on_disk = os.path.getsize(temp_path)
        use_stream = ins.should_use_streaming_ingest(on_disk)
        cat = category.strip() or "默认知识库"
        desc = description or ""

        # ---------- 流式入库（大文件）----------
        if use_stream and file_ext in (".txt", ".md"):
            enc = ins.detect_text_file_encoding(temp_path)
            sample = ins.read_text_head(temp_path, enc, 18000)
            summary = ins.make_rule_summary_from_sample(
                sample, uploaded_file.name, file_ext.lstrip(".")
            )
            n = ins.run_streaming_ingest(
                ins.iter_segments_text_file(temp_path, enc),
                uploaded_file.name,
                file_ext.lstrip("."),
                vector_db,
                on_disk,
                summary,
            )
            print(f"[Ingest/stream] txt/md 入库 {n} 块，文件：{uploaded_file.name}")
            _finalize_ingest_metadata(uploaded_file, file_ext, cat, desc, n)
            return n

        if use_stream and file_ext == ".pdf":
            if ins.pdf_has_extractable_text(temp_path):
                sample = ins.read_pdf_text_sample(temp_path)
                summary = ins.make_rule_summary_from_sample(sample, uploaded_file.name, "pdf")
                n = ins.run_streaming_ingest(
                    ins.iter_segments_pdf_text(temp_path),
                    uploaded_file.name,
                    "pdf",
                    vector_db,
                    on_disk,
                    summary,
                )
            else:
                n = ins.run_streaming_ingest(
                    ins.iter_segments_pdf_ocr(temp_path),
                    uploaded_file.name,
                    "pdf",
                    vector_db,
                    on_disk,
                    None,
                )
            print(f"[Ingest/stream] PDF 入库 {n} 块，文件：{uploaded_file.name}")
            _finalize_ingest_metadata(uploaded_file, file_ext, cat, desc, n)
            return n

        if use_stream and file_ext in (".docx", ".doc"):
            sample = ins.read_docx_sample(temp_path)
            summary = ins.make_rule_summary_from_sample(sample, uploaded_file.name, "docx")
            n = ins.run_streaming_ingest(
                ins.iter_segments_docx(temp_path),
                uploaded_file.name,
                "docx",
                vector_db,
                on_disk,
                summary,
            )
            print(f"[Ingest/stream] DOCX 入库 {n} 块，文件：{uploaded_file.name}")
            _finalize_ingest_metadata(uploaded_file, file_ext, cat, desc, n)
            return n

        # ---------- 原有路径（较小文件）：全文进内存 + 多层级 smart chunk ----------
        docs = []
        if file_ext == ".txt":
            with open(temp_path, "rb") as f:
                raw_data = f.read()

            decoded_text = None
            for enc in ["utf-8", "gb18030", "gbk", "latin-1"]:
                try:
                    decoded_text = raw_data.decode(enc)
                    break
                except UnicodeDecodeError:
                    continue

            if decoded_text is None:
                raise ValueError(f"无法解码文件 {uploaded_file.name}，不支持的编码格式")

            docs = [
                Document(
                    page_content=decoded_text.strip(),
                    metadata={"source_file": uploaded_file.name},
                )
            ]

        elif file_ext == ".pdf":
            try:
                loader = PyPDFLoader(temp_path)
                pdf_docs = loader.load()

                has_text = any(doc.page_content.strip() for doc in pdf_docs)

                if not has_text:
                    st_images = convert_from_path(
                        temp_path,
                        poppler_path=None,
                        fmt="png",
                        dpi=300,
                    )

                    ocr_text = []
                    for i, img in enumerate(st_images):
                        try:
                            page_text = pytesseract.image_to_string(img, lang="chi_sim")
                            ocr_text.append(f"第{i + 1}页：\n{page_text}")
                        except Exception as e:
                            print(f"第{i + 1}页OCR失败: {e}")
                            ocr_text.append(f"第{i + 1}页：OCR识别失败")

                    full_text = "\n\n".join(ocr_text)
                    docs = [
                        Document(
                            page_content=full_text,
                            metadata={"source_file": uploaded_file.name},
                        )
                    ]
                else:
                    docs = []
                    for i, doc in enumerate(pdf_docs):
                        doc.metadata.update({"source_file": uploaded_file.name, "page": i + 1})
                        docs.append(doc)

            except Exception as e:
                raise RuntimeError(f"PDF处理失败: {str(e)}\n{traceback.format_exc()}")

        elif file_ext in [".docx", ".doc"]:
            try:
                docx_file = DocxDocument(temp_path)
                paragraphs = []
                for para in docx_file.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text.strip())

                full_text = "\n\n".join(paragraphs)

                for table in docx_file.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        table_text.append(" | ".join(row_text))
                    if table_text:
                        full_text += "\n\n" + "\n".join(table_text)

                if not full_text.strip():
                    raise ValueError("DOCX文件中未提取到文本内容")

                docs = [
                    Document(
                        page_content=full_text,
                        metadata={"source_file": uploaded_file.name, "file_type": "docx"},
                    )
                ]
            except Exception as e:
                raise RuntimeError(f"DOCX处理失败: {str(e)}\n{traceback.format_exc()}")

        elif file_ext == ".md":
            try:
                with open(temp_path, "r", encoding="utf-8") as f:
                    md_text = f.read()

                if not md_text.strip():
                    raise ValueError("Markdown文件中未提取到文本内容")

                docs = [
                    Document(
                        page_content=md_text,
                        metadata={"source_file": uploaded_file.name, "file_type": "md"},
                    )
                ]
            except UnicodeDecodeError:
                with open(temp_path, "r", encoding="gb18030") as f:
                    md_text = f.read()
                docs = [
                    Document(
                        page_content=md_text,
                        metadata={"source_file": uploaded_file.name, "file_type": "md"},
                    )
                ]
            except Exception as e:
                raise RuntimeError(f"Markdown处理失败: {str(e)}\n{traceback.format_exc()}")

        elif file_ext in [".xlsx", ".xls"]:
            try:
                import pandas as pd

                excel_file = pd.ExcelFile(temp_path)
                sheets_text = []

                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    sheet_text = f"工作表: {sheet_name}\n\n"
                    sheet_text += df.to_string(index=False)
                    sheets_text.append(sheet_text)

                full_text = "\n\n" + "=" * 50 + "\n\n".join(sheets_text)

                if not full_text.strip():
                    raise ValueError("Excel文件中未提取到文本内容")

                docs = [
                    Document(
                        page_content=full_text,
                        metadata={"source_file": uploaded_file.name, "file_type": "excel"},
                    )
                ]
            except ImportError:
                raise RuntimeError("处理Excel文件需要安装pandas和openpyxl库，请运行: pip install pandas openpyxl")
            except Exception as e:
                raise RuntimeError(f"Excel处理失败: {str(e)}\n{traceback.format_exc()}")

        else:
            print(f"不支持的文件类型: {file_ext}")
            return 0

        from utils.smart_chunker import smart_chunk_document

        full_text = "\n\n".join([doc.page_content for doc in docs])
        text_length = len(full_text)
        if text_length < 5000:
            length_factor = 0.8
        elif text_length > 50000:
            length_factor = 1.2
        else:
            length_factor = 1.0

        chunks, chunk_stats = smart_chunk_document(
            text=full_text,
            source_file=uploaded_file.name,
            file_type=file_ext.lstrip("."),
            use_llm_summary=False,
            doc_length_factor=length_factor,
        )

        print(f"[SmartChunker] 分块统计: {chunk_stats}")

        for chunk in chunks:
            if "source_file" not in chunk.metadata:
                chunk.metadata["source_file"] = uploaded_file.name
            if "file_type" not in chunk.metadata:
                chunk.metadata["file_type"] = file_ext.lstrip(".")

        if chunks:
            index_dir = os.path.join(get_kb_dir(), "faiss_index")
            os.makedirs(index_dir, exist_ok=True)
            bs = ins.EMBED_ADD_BATCH_SIZE
            for i in range(0, len(chunks), bs):
                vector_db.add_documents(chunks[i : i + bs])
            vector_db.save_local(index_dir)
            print(f"成功入库 {len(chunks)} 个文本块，文件：{uploaded_file.name}")
            _finalize_ingest_metadata(uploaded_file, file_ext, cat, desc, len(chunks))

        return len(chunks)

    except Exception as e:
        error_msg = f"文件处理失败 {uploaded_file.name}: {str(e)}"
        print(error_msg)
        print(traceback.format_exc())
        log_error("file_processing", error_msg, {"file_name": uploaded_file.name})
        raise

    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except Exception as e:
                print(f"删除临时文件失败 {temp_path}: {e}")
