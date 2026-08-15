"""即时文档：从上传字节解析纯文本，供 Web API 使用（不入库、不经向量库）。"""
from __future__ import annotations

import os
import tempfile
from typing import List

from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document


INSTANT_MAX_UPLOAD_BYTES = 5 * 1024 * 1024
INSTANT_MAX_TEXT_CHARS = 100_000
INSTANT_ALLOWED_EXT = {".pdf", ".txt", ".md", ".docx"}


def _read_txt_md(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(path, "r", encoding="gbk") as f:
            return f.read()


def parse_upload_bytes(filename: str, data: bytes) -> str:
    """
    解析上传文件为单一字符串。超长文本在调用方截断并报错。
    支持：pdf / txt / md / docx。
    """
    raw = filename or "unnamed"
    ext = os.path.splitext(raw.replace("\\", "/"))[1].lower()
    if ext not in INSTANT_ALLOWED_EXT:
        raise ValueError(f"不支持的格式，仅允许：{', '.join(sorted(INSTANT_ALLOWED_EXT))}")
    if len(data) > INSTANT_MAX_UPLOAD_BYTES:
        raise ValueError(f"文件超过 {INSTANT_MAX_UPLOAD_BYTES // (1024 * 1024)}MB 限制")

    temp_path: str | None = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext, mode="wb") as tmp:
            tmp.write(data)
            temp_path = tmp.name

        docs: List[Document] = []
        if ext == ".pdf":
            loader = PyPDFLoader(temp_path)
            docs = loader.load()
        elif ext in (".txt", ".md"):
            content = _read_txt_md(temp_path)
            if not content.strip():
                raise ValueError("文本文件为空")
            docs = [Document(page_content=content, metadata={"source_file": raw})]
        elif ext == ".docx":
            docx_doc = DocxDocument(temp_path)
            paragraphs: List[str] = []
            for para in docx_doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            for table in docx_doc.tables:
                rows_txt: List[str] = []
                for row in table.rows:
                    rows_txt.append(" | ".join(cell.text.strip() for cell in row.cells))
                if rows_txt:
                    paragraphs.append("\n".join(rows_txt))
            full_text = "\n\n".join(paragraphs)
            if not full_text.strip():
                raise ValueError("Word 文档中未提取到文本")
            docs = [Document(page_content=full_text, metadata={"source_file": raw})]

        text = "\n\n".join(d.page_content for d in docs if d.page_content)
        text = text.strip()
        if not text:
            raise ValueError("未能从文件中提取文本")
        if len(text) > INSTANT_MAX_TEXT_CHARS:
            raise ValueError(f"提取的正文超过 {INSTANT_MAX_TEXT_CHARS:,} 个字符，请缩短或拆分文档")
        return text
    finally:
        if temp_path and os.path.isfile(temp_path):
            try:
                os.unlink(temp_path)
            except OSError:
                pass
