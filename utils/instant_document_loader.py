# utils/instant_document_loader.py
"""
即时文档加载器：用于文档问答页面
不经过向量库，直接解析文档内容供LLM使用
解析方法与ChatGPT、DeepSeek、豆包一致
"""
import os
import tempfile
import traceback
from typing import List
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument
import streamlit as st


def parse_document_instant(uploaded_file) -> List[Document]:
    """
    即时解析文档（不经过向量库）
    解析方法与ChatGPT、DeepSeek、豆包一致
    
    :param uploaded_file: Streamlit上传的文件对象
    :return: 解析后的文档列表
    """
    temp_path = None
    try:
        # 1. 创建临时文件
        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        with tempfile.NamedTemporaryFile(
                delete=False,
                suffix=file_ext,
                mode='wb'
        ) as tmp:
            tmp.write(uploaded_file.getbuffer())
            temp_path = tmp.name

        # 2. 根据文件类型解析
        docs = []
        
        if file_ext == '.pdf':
            # PDF文件：使用PyPDFLoader（与ChatGPT等方法一致）
            try:
                loader = PyPDFLoader(temp_path)
                docs = loader.load()
                # 为每个文档添加文件名元数据
                for doc in docs:
                    doc.metadata["source_file"] = uploaded_file.name
                    doc.metadata["file_type"] = "pdf"
            except Exception as e:
                raise RuntimeError(f"PDF解析失败: {str(e)}")

        elif file_ext in ['.txt', '.md']:
            # 文本文件：直接读取
            try:
                with open(temp_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                if not content.strip():
                    raise ValueError("文本文件为空")
                docs = [Document(
                    page_content=content,
                    metadata={"source_file": uploaded_file.name, "file_type": file_ext.lstrip('.')}
                )]
            except UnicodeDecodeError:
                # 尝试其他编码
                with open(temp_path, 'r', encoding='gbk') as f:
                    content = f.read()
                docs = [Document(
                    page_content=content,
                    metadata={"source_file": uploaded_file.name, "file_type": file_ext.lstrip('.')}
                )]
            except Exception as e:
                raise RuntimeError(f"文本文件解析失败: {str(e)}")

        elif file_ext == '.docx':
            # Word文档（.docx）：使用python-docx
            try:
                docx_doc = DocxDocument(temp_path)
                paragraphs = []
                for para in docx_doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                # 提取表格内容
                for table in docx_doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        table_text.append(row_text)
                    if table_text:
                        paragraphs.append("\n".join(table_text))
                
                full_text = "\n\n".join(paragraphs)
                if not full_text.strip():
                    raise ValueError("Word文档中未提取到文本内容")
                
                docs = [Document(
                    page_content=full_text,
                    metadata={"source_file": uploaded_file.name, "file_type": "docx"}
                )]
            except Exception as e:
                raise RuntimeError(f"Word文档解析失败: {str(e)}")
        
        elif file_ext == '.doc':
            # Word文档（.doc）：使用docx2txt或textract
            try:
                # 尝试使用docx2txt（需要安装：pip install docx2txt）
                try:
                    import docx2txt
                    full_text = docx2txt.process(temp_path)
                    if not full_text or not full_text.strip():
                        raise ValueError("DOC文件中未提取到文本内容")
                    
                    docs = [Document(
                        page_content=full_text.strip(),
                        metadata={"source_file": uploaded_file.name, "file_type": "doc"}
                    )]
                except ImportError:
                    # 如果docx2txt不可用，尝试使用textract
                    try:
                        import textract
                        full_text = textract.process(temp_path).decode('utf-8')
                        if not full_text or not full_text.strip():
                            raise ValueError("DOC文件中未提取到文本内容")
                        
                        docs = [Document(
                            page_content=full_text.strip(),
                            metadata={"source_file": uploaded_file.name, "file_type": "doc"}
                        )]
                    except ImportError:
                        raise RuntimeError(
                            "解析.doc文件需要安装docx2txt或textract库。\n"
                            "请运行: pip install docx2txt\n"
                            "或者: pip install textract"
                        )
            except Exception as e:
                raise RuntimeError(f"DOC文件解析失败: {str(e)}")

        elif file_ext in ['.xlsx', '.xls']:
            # Excel文件：使用pandas
            try:
                import pandas as pd
                excel_file = pd.ExcelFile(temp_path)
                sheets_text = []
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    sheet_text = f"工作表: {sheet_name}\n\n"
                    # 转换为文本格式
                    sheet_text += df.to_string(index=False)
                    sheets_text.append(sheet_text)
                
                full_text = "\n\n" + "="*50 + "\n\n".join(sheets_text)
                
                if not full_text.strip():
                    raise ValueError("Excel文件中未提取到文本内容")
                
                docs = [Document(
                    page_content=full_text,
                    metadata={"source_file": uploaded_file.name, "file_type": "excel"}
                )]
            except ImportError:
                raise RuntimeError("处理Excel文件需要安装pandas和openpyxl库")
            except Exception as e:
                raise RuntimeError(f"Excel解析失败: {str(e)}")

        else:
            raise ValueError(f"不支持的文件类型: {file_ext}")

        return docs

    except Exception as e:
        st.error(f"文档解析失败: {str(e)}")
        return []
    finally:
        # 清理临时文件
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except:
                pass


def combine_documents(docs_list: List[List[Document]]) -> str:
    """
    合并多个文档的内容为单个文本字符串
    
    :param docs_list: 文档列表的列表（每个文件对应一个文档列表）
    :return: 合并后的文本内容
    """
    combined_parts = []
    
    for file_docs in docs_list:
        if not file_docs:
            continue
        
        # 获取文件名
        file_name = file_docs[0].metadata.get("source_file", "未知文件")
        file_type = file_docs[0].metadata.get("file_type", "")
        
        # 合并该文件的所有页面/段落
        file_content = "\n\n".join([doc.page_content for doc in file_docs])
        
        # 添加文件标识
        combined_parts.append(f"【文件：{file_name}】\n{file_content}")
    
    return "\n\n" + "="*80 + "\n\n".join(combined_parts)

