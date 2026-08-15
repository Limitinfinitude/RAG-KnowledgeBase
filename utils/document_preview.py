# utils/document_preview.py
"""
文档预览 / 全文查看：优先从入库时保存的原文副本解析；若无副本则从向量块按顺序拼接。
"""
import logging
import os
import shutil
import unicodedata
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument

from utils.metadata_manager import get_document_metadata
from utils.path_context import get_kb_dir

ORIGINAL_FILES_SUBDIR = "original_files"
MAX_FULL_VIEW_CHARS = 500_000


def _original_abs_path(file_name: str) -> str:
    base = os.path.basename((file_name or "").replace("\\", "/"))
    return os.path.join(get_kb_dir(), ORIGINAL_FILES_SUBDIR, base)


def persist_original_from_temp(temp_path: str, logical_name: str) -> None:
    """入库时将临时文件复制为原文存档，供「查看内容」使用。"""
    if not temp_path or not os.path.isfile(temp_path):
        return
    base = os.path.basename((logical_name or "file").replace("\\", "/"))
    if not base or base in (".", ".."):
        return
    try:
        dest_dir = os.path.join(get_kb_dir(), ORIGINAL_FILES_SUBDIR)
        os.makedirs(dest_dir, exist_ok=True)
        dest = os.path.join(dest_dir, base)
        # 同名校验后覆盖，与元数据「按文件名索引」一致
        shutil.copy2(temp_path, dest)
    except OSError as e:
        logger.warning("[Ingest] 保存原文副本失败 %s: %s", logical_name, e)


def _sort_docs_for_reconstruct(file_name: str, file_docs: List[Any]) -> List[Any]:
    def key_fn(d: Any) -> Tuple:
        md = getattr(d, "metadata", None) or {}
        page = md.get("page", 0)
        try:
            page = int(page)
        except (TypeError, ValueError):
            page = 0
        cs = md.get("chunk_start")
        try:
            cs_i = int(cs) if cs is not None else 0
        except (TypeError, ValueError):
            cs_i = 0
        cid = str(md.get("chunk_id") or "")
        return (page, cs_i, cid)

    out = list(file_docs)
    if (file_name or "").lower().endswith(".pdf"):
        out.sort(key=key_fn)
    else:
        out.sort(key=key_fn)
    return out


def reconstruct_full_text_from_vector_db(file_name: str, vector_db) -> str:
    """无原文存档时，按元数据顺序拼接全部块（非二进制意义上的原文件）。"""
    docs = vector_db.similarity_search("", k=30000)
    file_docs = [d for d in docs if d.metadata.get("source_file") == file_name]
    if not file_docs:
        return ""
    file_docs = _sort_docs_for_reconstruct(file_name, file_docs)
    parts = [d.page_content for d in file_docs if getattr(d, "page_content", None)]
    return "\n\n".join(parts)


def _read_txt_like(path: str) -> str:
    for enc in ("utf-8", "gb18030", "gbk", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(path, "rb") as f:
        return f.read().decode("utf-8", errors="replace")


def extract_plain_text_from_original_path(abs_path: str, *, max_chars: int) -> str:
    """从本地原文文件提取纯文本（知识库查看用，限制最大字符）。"""
    ext = os.path.splitext(abs_path)[1].lower()
    text = ""
    if ext in (".txt", ".md"):
        text = _read_txt_like(abs_path)
    elif ext == ".pdf":
        loader = PyPDFLoader(abs_path)
        pdf_docs = loader.load()
        text = "\n\n".join(d.page_content for d in pdf_docs if d.page_content)
    elif ext in (".docx", ".doc"):
        docx_doc = DocxDocument(abs_path)
        paragraphs: List[str] = []
        for para in docx_doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        for table in docx_doc.tables:
            rows_txt: List[str] = []
            for row in table.rows:
                rows_txt.append(" | ".join(cell.text.strip() for cell in row.cells))
            if rows_txt:
                paragraphs.append("\n".join(rows_txt))
        text = "\n\n".join(paragraphs)
    elif ext in (".xlsx", ".xls"):
        try:
            import pandas as pd

            excel_file = pd.ExcelFile(abs_path)
            sheets_text: List[str] = []
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                sheet_text = f"工作表: {sheet_name}\n\n"
                sheet_text += df.to_string(index=False)
                sheets_text.append(sheet_text)
            text = "\n\n" + "=" * 50 + "\n\n".join(sheets_text)
        except ImportError:
            raise RuntimeError("查看 Excel 需要 pandas / openpyxl，请安装依赖") from None
    else:
        raise ValueError(f"暂不支持从该扩展名提取全文：{ext}")

    text = (text or "").strip()
    if len(text) > max_chars:
        return text[:max_chars]
    return text


def get_document_full_view_payload(
    file_name: str,
    vector_db,
    *,
    max_chars: int = MAX_FULL_VIEW_CHARS,
) -> Dict[str, Any]:
    """
    供 API 返回：整篇纯文本视图（非分块列表）。
    source: original_file | reconstructed
    """
    orig = _original_abs_path(file_name)
    truncated = False
    text = ""
    source = "reconstructed"
    err: Optional[str] = None

    if os.path.isfile(orig):
        try:
            text = extract_plain_text_from_original_path(orig, max_chars=max_chars)
            source = "original_file"
            if len(text) >= max_chars:
                truncated = True
        except Exception as e:
            err = str(e)
            text = ""
            source = "reconstructed"

    if not text.strip():
        try:
            text = reconstruct_full_text_from_vector_db(file_name, vector_db)
            source = "reconstructed"
        except Exception as e:
            return {
                "file_name": file_name,
                "text": "",
                "source": "error",
                "truncated": False,
                "error": str(e),
            }
        if len(text) > max_chars:
            text = text[:max_chars]
            truncated = True

    if err and source == "reconstructed" and not text.strip():
        return {
            "file_name": file_name,
            "text": "",
            "source": "error",
            "truncated": False,
            "error": err or "无法提取正文",
        }

    if not (text or "").strip():
        return {
            "file_name": file_name,
            "text": "",
            "source": "error",
            "truncated": False,
            "error": "未找到该文档或正文中无可用内容",
        }

    return {
        "file_name": file_name,
        "text": text,
        "source": source,
        "truncated": truncated,
        "note": (
            "正文由入库时保存的源文件解析。"
            if source == "original_file"
            else "未找到本地原文副本（可能为旧数据），以下为按分块顺序拼接的全文，仅供参考。"
        ),
        "parse_warning": err,
    }


def get_plain_text_for_kb_substring_search(
    file_name: str,
    vector_db,
    *,
    max_chars: int = 1_500_000,
) -> str:
    """加载用于子串检索的纯文本（与「查看内容」同源，不调用向量相似度）。"""
    pl = get_document_full_view_payload(file_name, vector_db, max_chars=max_chars)
    if pl.get("source") == "error":
        return ""
    return (pl.get("text") or "").strip()


def _normalize_search_text(s: str) -> str:
    """统一 Unicode 形式，减少因兼容字符导致的漏检/误检。"""
    return unicodedata.normalize("NFC", s or "")


def _find_casefold_substring(haystack: str, query: str, start: int) -> int:
    """
    在 haystack 中找与 query 大小写折叠相等的子串起点（基于原文下标切片）。
    仅当「全文与关键词 casefold 后字符数与原文一致」时用 fast find，否则逐位比较，
    避免 ß→ss、某些 Unicode 规则导致「在 casefold 串上 find 再映射回原文」错位。
    """
    h = haystack
    q = query
    if not q or start >= len(h):
        return -1
    q_cf = q.casefold()
    q_len = len(q)
    if start > len(h) - q_len:
        return -1
    h_cf = h.casefold()
    if len(h_cf) == len(h) and len(q_cf) == len(q):
        idx = h_cf.find(q_cf, start)
        if idx < 0:
            return -1
        if h[idx : idx + q_len].casefold() != q_cf:
            return -1
        return idx
    i = start
    limit = len(h) - q_len + 1
    while i < limit:
        if h[i : i + q_len].casefold() == q_cf:
            return i
        i += 1
    return -1


def substring_hits_with_context(
    text: str,
    query: str,
    *,
    context_before: int = 90,
    context_after: int = 120,
    max_hits: int = 500,
) -> List[Dict[str, Any]]:
    """
    全文子串命中（大小写不敏感、非重叠），每条含原文中的 before / match / after 与 global_offset，供前端高亮。
    索引始终基于原文切片，避免 casefold 扩长字符导致 find 偏移错位。
    """
    raw = _normalize_search_text(text)
    q = _normalize_search_text((query or "").strip())
    out: List[Dict[str, Any]] = []
    if not q or not raw:
        return out
    q_cf = q.casefold()
    qlen = len(q)
    pos = 0
    while len(out) < max_hits:
        idx = _find_casefold_substring(raw, q, pos)
        if idx < 0:
            break
        match_s = raw[idx : idx + qlen]
        if match_s.casefold() != q_cf:
            pos = idx + 1
            continue
        mb = max(0, idx - context_before)
        me = min(len(raw), idx + qlen + context_after)
        before = raw[mb:idx]
        after = raw[idx + qlen : me]
        if mb > 0:
            before = "…" + before
        if me < len(raw):
            after = after + "…"
        out.append(
            {
                "global_offset": idx,
                "before": before,
                "match": match_s,
                "after": after,
            }
        )
        pos = idx + qlen
    return out


def snippet_for_substring_match(full_text: str, query: str, *, radius_chars: int = 140) -> str:
    """在正文中查找 query 的首次出现（大小写不敏感），截取前后片段供列表展示。"""
    hits = substring_hits_with_context(
        full_text, query, context_before=radius_chars, context_after=radius_chars, max_hits=1
    )
    if hits:
        h = hits[0]
        return f"{h['before']}{h['match']}{h['after']}"
    text = full_text or ""
    return text[:320] + ("…" if len(text) > 320 else "")


def get_document_structure(file_name: str, vector_db) -> Dict:
    """
    获取文档结构信息
    :param file_name: 文件名
    :param vector_db: 向量数据库实例
    :return: 文档结构信息
    """
    try:
        # 从向量库中获取该文档的所有chunks
        docs = vector_db.similarity_search("", k=30000)
        file_docs = [d for d in docs if d.metadata.get("source_file") == file_name]

        if not file_docs:
            return {
                "file_name": file_name,
                "chunks_count": 0,
                "total_chars": 0,
                "pages": [],
                "structure": "未找到文档",
            }

        # 统计信息
        chunks_count = len(file_docs)
        total_chars = sum(len(d.page_content) for d in file_docs)

        # 获取页面信息（如果是PDF）
        pages = []
        if file_name.lower().endswith(".pdf"):
            for doc in file_docs:
                page_num = doc.metadata.get("page", 0)
                if page_num > 0:
                    pages.append(
                        {
                            "page": page_num,
                            "chars": len(doc.page_content),
                            "preview": doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content,
                        }
                    )

        # 获取元数据
        metadata = get_document_metadata(file_name)

        return {
            "file_name": file_name,
            "chunks_count": chunks_count,
            "total_chars": total_chars,
            "pages": sorted(pages, key=lambda x: x["page"]) if pages else [],
            "metadata": metadata,
            "structure": "PDF文档" if file_name.lower().endswith(".pdf") else "文档",
            "has_original_file": os.path.isfile(_original_abs_path(file_name)),
        }
    except Exception as e:
        return {
            "file_name": file_name,
            "error": str(e),
        }


def preview_document_content(file_name: str, vector_db, max_chunks: int = 10) -> List[Dict]:
    """
    预览文档内容（显示前N个chunks）— 保留兼容；新 UI 请用 get_document_full_view_payload。
    """
    try:
        docs = vector_db.similarity_search("", k=30000)
        file_docs = [d for d in docs if d.metadata.get("source_file") == file_name]

        if not file_docs:
            return []

        # 按页面或顺序排序
        if file_name.lower().endswith(".pdf"):
            file_docs.sort(key=lambda x: x.metadata.get("page", 0))

        preview_docs = file_docs[:max_chunks]

        result = []
        for i, doc in enumerate(preview_docs):
            result.append(
                {
                    "chunk_id": i + 1,
                    "page": doc.metadata.get("page", "N/A"),
                    "content": doc.page_content,
                    "chars": len(doc.page_content),
                }
            )

        return result
    except Exception as e:
        return [{"error": str(e)}]


def get_file_type_icon(file_type: str) -> str:
    """根据文件类型返回图标"""
    icons = {
        "pdf": "📄",
        "txt": "📝",
        "docx": "📘",
        "doc": "📘",
        "md": "📋",
        "xlsx": "📊",
        "xls": "📊",
    }
    return icons.get(file_type.lower(), "📄")
